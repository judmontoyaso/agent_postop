"""scripts/build_docx.py — Genera el informe final en .docx desde el Markdown.

El entregable se pide en un formato de documento, pero la fuente de verdad
sigue siendo `docs/final-report.md`: se versiona, se revisa en el diff y no se
desincroniza del repositorio. Este script lo convierte, así que actualizar el
Word es volver a ejecutarlo en vez de reeditar a mano.

Convierte lo que el informe usa —encabezados, listas, tablas, citas, imágenes,
negritas, cursivas y código en línea— y no pretende ser un conversor Markdown
completo.

Uso:  python scripts/build_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
ORIGEN = RAIZ / "docs" / "final-report.md"
DESTINO = RAIZ / "docs" / "Informe final - Agente de voz postoperatorio.docx"

GRIS = RGBColor(0x55, 0x55, 0x55)


def aplicar_formato(parrafo, texto: str) -> None:
    """Interpreta **negrita**, *cursiva* y `código` dentro de una línea."""
    for trozo in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", texto):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            parrafo.add_run(trozo[2:-2]).bold = True
        elif trozo.startswith("*") and trozo.endswith("*"):
            parrafo.add_run(trozo[1:-1]).italic = True
        elif trozo.startswith("`") and trozo.endswith("`"):
            run = parrafo.add_run(trozo[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            parrafo.add_run(trozo)


def añadir_tabla(doc: Document, filas: list[str]) -> None:
    """Las tablas Markdown traen una fila de guiones que no es contenido."""
    celdas = [[c.strip() for c in f.strip().strip("|").split("|")] for f in filas]
    cuerpo = [f for f in celdas if not all(set(c) <= set("-: ") for c in f)]
    if not cuerpo:
        return

    tabla = doc.add_table(rows=len(cuerpo), cols=len(cuerpo[0]))
    tabla.style = "Light Grid Accent 1"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, fila in enumerate(cuerpo):
        for j, valor in enumerate(fila):
            if j >= len(tabla.rows[i].cells):
                continue
            celda = tabla.rows[i].cells[j]
            celda.text = ""
            aplicar_formato(celda.paragraphs[0], valor)
            if i == 0:
                for run in celda.paragraphs[0].runs:
                    run.bold = True


def convertir() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lineas = ORIGEN.read_text(encoding="utf-8").split("\n")
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        pelada = linea.strip()

        if not pelada or pelada == "---":
            i += 1
            continue

        # Imagen: ![alt](archivo.png)
        img = re.match(r"!\[(.*?)\]\((.+?)\)", pelada)
        if img:
            ruta = ORIGEN.parent / img.group(2)
            if ruta.exists():
                doc.add_picture(str(ruta), width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                pie = doc.add_paragraph()
                pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pie.add_run(img.group(1))
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = GRIS
            i += 1
            continue

        # Tabla: se consumen todas sus filas de golpe
        if pelada.startswith("|"):
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(lineas[i])
                i += 1
            añadir_tabla(doc, filas)
            doc.add_paragraph()
            continue

        if pelada.startswith("#"):
            nivel = len(pelada) - len(pelada.lstrip("#"))
            doc.add_heading(pelada.lstrip("# ").strip(), level=min(nivel, 4))
            i += 1
            continue

        if pelada.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            aplicar_formato(p, pelada.lstrip("> ").strip())
            i += 1
            continue

        if re.match(r"^[-*]\s+", pelada):
            p = doc.add_paragraph(style="List Bullet")
            aplicar_formato(p, re.sub(r"^[-*]\s+", "", pelada))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", pelada):
            p = doc.add_paragraph(style="List Number")
            aplicar_formato(p, re.sub(r"^\d+\.\s+", "", pelada))
            i += 1
            continue

        # Párrafo: el Markdown parte las líneas a 79 columnas, así que se unen
        # hasta la línea en blanco para que el Word no quede con saltos raros.
        bloque = []
        while i < len(lineas) and lineas[i].strip() and not re.match(
            r"^\s*([#>|]|[-*]\s|\d+\.\s|!\[)", lineas[i]
        ):
            bloque.append(lineas[i].strip())
            i += 1
        if bloque:
            aplicar_formato(doc.add_paragraph(), " ".join(bloque))

    doc.save(DESTINO)
    print(f"Generado: {DESTINO.name}")
    print(f"  {DESTINO.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    convertir()
